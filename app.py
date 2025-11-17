import streamlit as st
import io
from docx import Document
from docx.shared import Pt, RGBColor


PRELOADED_PROGRAMS = {
"Program 1: Linear and Binary Search": """#include <iostream>
#include <algorithm> // For std::sort

using namespace std;

// Linear search: O(n)
int linearSearch(int arr[], int n, int key) {
    for (int i = 0; i < n; i++) {
        if (arr[i] == key) {
            return i; // Return index if found
        }
    }
    return -1; // Return -1 if not found
}

// Binary search: O(log n)
// NOTE: Requires the array to be sorted first.
int binarySearch(int arr[], int n, int key) {
    int low = 0, high = n - 1;
    while (low <= high) {
        int mid = low + (high - low) / 2; // Better way to avoid overflow
        if (arr[mid] == key) {
            return mid; // Return index if found
        } else if (arr[mid] < key) {
            low = mid + 1;
        } else {
            high = mid - 1;
        }
    }
    return -1; // Return -1 if not found
}

void printArray(int arr[], int n) {
    for (int i = 0; i < n; i++) {
        cout << arr[i] << " ";
    }
    cout << endl;
}

int main() {
    int n;
    cout << "Enter number of elements: ";
    cin >> n;
    int* arr = new int[n];
    cout << "Enter " << n << " elements: " << endl;
    for(int i = 0; i < n; i++) {
        cin >> arr[i];
    }

    int key;
    cout << "Enter key to search for: ";
    cin >> key;

    // Linear Search
    int linearResult = linearSearch(arr, n, key);
    if (linearResult != -1) {
        cout << "Linear Search: Found at index " << linearResult << endl;
    } else {
        cout << "Linear Search: Key not found." << endl;
    }

    // Binary Search
    // We must sort the array first!
    sort(arr, arr + n);
    cout << "Array sorted for Binary Search: ";
    printArray(arr, n);
    
    int binaryResult = binarySearch(arr, n, key);
    if (binaryResult != -1) {
        cout << "Binary Search: Found at index " << binaryResult << endl;
    } else {
        cout << "Binary Search: Key not found." << endl;
    }

    delete[] arr;
    return 0;
}""",

"Program 2: Stack using Array": """#include <iostream>
#include <climits> // For INT_MIN

using namespace std;

#define MAX_SIZE 100

class Stack {
    int top;
    int arr[MAX_SIZE];

public:
    Stack() { top = -1; }

    bool isEmpty() {
        return (top == -1);
    }

    bool isFull() {
        return (top == MAX_SIZE - 1);
    }

    void push(int x) {
        if (isFull()) {
            cout << "Error: Stack Overflow" << endl;
            return;
        }
        arr[++top] = x;
        cout << x << " pushed to stack" << endl;
    }

    int pop() {
        if (isEmpty()) {
            cout << "Error: Stack Underflow" << endl;
            return INT_MIN;
        }
        int x = arr[top--];
        return x;
    }

    int peek() {
        if (isEmpty()) {
            cout << "Error: Stack is Empty" << endl;
            return INT_MIN;
        }
        int x = arr[top];
        return x;
    }
};

int main() {
    Stack s;
    int choice, val;

    do {
        cout << "\n--- Stack Menu ---\n";
        cout << "1. Push\n";
        cout << "2. Pop\n";
        cout << "3. Peek\n";
        cout << "4. Check if Empty\n";
        cout << "5. Check if Full\n";
        cout << "0. Exit\n";
        cout << "Enter choice: ";
        cin >> choice;

        switch (choice) {
            case 1:
                cout << "Enter value to push: ";
                cin >> val;
                s.push(val);
                break;
            case 2:
                val = s.pop();
                if (val != INT_MIN) {
                    cout << val << " popped from stack" << endl;
                }
                break;
            case 3:
                val = s.peek();
                if (val != INT_MIN) {
                    cout << "Top element is " << val << endl;
                }
                break;
            case 4:
                cout << (s.isEmpty() ? "Stack is empty" : "Stack is not empty") << endl;
                break;
            case 5:
                cout << (s.isFull() ? "Stack is full" : "Stack is not full") << endl;
                break;
            case 0:
                cout << "Exiting..." << endl;
                break;
            default:
                cout << "Invalid choice." << endl;
        }
    } while (choice != 0);

    return 0;
}""",

"Program 3: Queue using Array (Circular)": """#include <iostream>
#include <climits> // For INT_MIN

using namespace std;

#define MAX_SIZE 100

class Queue {
    int front, rear, size;
    int arr[MAX_SIZE];

public:
    Queue() {
        front = -1;
        rear = -1;
        size = 0;
    }

    bool isEmpty() {
        return (front == -1);
    }

    bool isFull() {
        return ((rear + 1) % MAX_SIZE == front);
    }

    void enqueue(int x) {
        if (isFull()) {
            cout << "Error: Queue Overflow" << endl;
            return;
        }
        if (isEmpty()) {
            front = 0; // Set front when enqueuing first element
        }
        rear = (rear + 1) % MAX_SIZE;
        arr[rear] = x;
        size++;
        cout << x << " enqueued to queue" << endl;
    }

    int dequeue() {
        if (isEmpty()) {
            cout << "Error: Queue Underflow" << endl;
            return INT_MIN;
        }
        int x = arr[front];
        if (front == rear) {
            // Last element is being dequeued
            front = -1;
            rear = -1;
        } else {
            front = (front + 1) % MAX_SIZE;
        }
        size--;
        return x;
    }

    int peek() {
        if (isEmpty()) {
            cout << "Error: Queue is Empty" << endl;
            return INT_MIN;
        }
        return arr[front];
    }
    
    int queueSize() {
        return size;
    }
};

int main() {
    Queue q;
    int choice, val;

    do {
        cout << "\n--- Queue Menu ---\n";
        cout << "1. Enqueue\n";
        cout << "2. Dequeue\n";
        cout << "3. Peek\n";
        cout << "4. Check if Empty\n";
        cout << "5. Check if Full\n";
        cout << "6. Get Size\n";
        cout << "0. Exit\n";
        cout << "Enter choice: ";
        cin >> choice;

        switch (choice) {
            case 1:
                cout << "Enter value to enqueue: ";
                cin >> val;
                q.enqueue(val);
                break;
            case 2:
                val = q.dequeue();
                if (val != INT_MIN) {
                    cout << val << " dequeued from queue" << endl;
                }
                break;
            case 3:
                val = q.peek();
                if (val != INT_MIN) {
                    cout << "Front element is " << val << endl;
                }
                break;
            case 4:
                cout << (q.isEmpty() ? "Queue is empty" : "Queue is not empty") << endl;
                break;
            case 5:
                cout << (q.isFull() ? "Queue is full" : "Queue is not full") << endl;
                break;
            case 6:
                cout << "Queue size is: " << q.queueSize() << endl;
                break;
            case 0:
                cout << "Exiting..." << endl;
                break;
            default:
                cout << "Invalid choice." << endl;
        }
    } while (choice != 0);

    return 0;
}""",

"Program 4: Linked List Operations": """#include <iostream>
using namespace std;

class Node {
public:
    int data;
    Node* next;
    Node(int d) {
        data = d;
        next = NULL;
    }
};

class List {
public:
    Node* head;
    List() { head = NULL; }

    bool isEmpty() {
        return head == NULL;
    }

    void insertBeg(int x) {
        Node* n = new Node(x);
        n->next = head;
        head = n;
        cout << x << " inserted at beginning" << endl;
    }

    void insertEnd(int x) {
        Node* n = new Node(x);
        if (isEmpty()) {
            head = n;
            cout << x << " inserted at end" << endl;
            return;
        }
        Node* t = head;
        while (t->next) {
            t = t->next;
        }
        t->next = n;
        cout << x << " inserted at end" << endl;
    }

    void insertAfter(int val, int x) {
        Node* t = head;
        while (t && t->data != val) {
            t = t->next;
        }
        if (!t) {
            cout << "Error: Value " << val << " not found." << endl;
            return;
        }
        Node* n = new Node(x);
        n->next = t->next;
        t->next = n;
        cout << x << " inserted after " << val << endl;
    }

    void deleteBeg() {
        if (isEmpty()) {
            cout << "Error: List is empty" << endl;
            return;
        }
        Node* t = head;
        head = head->next;
        cout << t->data << " deleted from beginning" << endl;
        delete t;
    }

    void deleteEnd() {
        if (isEmpty()) {
            cout << "Error: List is empty" << endl;
            return;
        }
        if (!head->next) { // Only one node
            cout << head->data << " deleted from end" << endl;
            delete head;
            head = NULL;
            return;
        }
        Node* t = head;
        while (t->next->next) {
            t = t->next;
        }
        cout << t->next->data << " deleted from end" << endl;
        delete t->next;
        t->next = NULL;
    }

    void deleteVal(int x) {
        if (isEmpty()) {
            cout << "Error: List is empty" << endl;
            return;
        }
        if (head->data == x) {
            deleteBeg();
            return;
        }
        Node* t = head;
        while (t->next && t->next->data != x) {
            t = t->next;
        }
        if (!t->next) {
            cout << "Error: Value " << x << " not found." << endl;
            return;
        }
        Node* d = t->next;
        t->next = d->next;
        cout << d->data << " deleted" << endl;
        delete d;
    }

    void display() {
        if (isEmpty()) {
            cout << "List is empty." << endl;
            return;
        }
        Node* t = head;
        cout << "List: HEAD -> ";
        while(t) {
            cout << t->data << " -> ";
            t = t->next;
        }
        cout << "NULL" << endl;
    }
};

int main() {
    List l;
    int choice, val, key;
    
    do {
        cout << "\n--- Linked List Menu ---\n";
        cout << "1. Insert at Beginning\n";
        cout << "2. Insert at End\n";
        cout << "3. Insert After Value\n";
        cout << "4. Delete from Beginning\n";
        cout << "5. Delete from End\n";
        cout << "6. Delete Specific Value\n";
        cout << "7. Display List\n";
        cout << "8. Check if Empty\n";
        cout << "0. Exit\n";
        cout << "Enter choice: ";
        cin >> choice;

        switch (choice) {
            case 1:
                cout << "Enter value: "; cin >> val;
                l.insertBeg(val);
                break;
            case 2:
                cout << "Enter value: "; cin >> val;
                l.insertEnd(val);
                break;
            case 3:
                cout << "Enter value to insert: "; cin >> val;
                cout << "Enter value to insert after: "; cin >> key;
                l.insertAfter(key, val);
                break;
            case 4:
                l.deleteBeg();
                break;
            case 5:
                l.deleteEnd();
                break;
            case 6:
                cout << "Enter value to delete: "; cin >> val;
                l.deleteVal(val);
                break;
            case 7:
                l.display();
                break;
            case 8:
                cout << (l.isEmpty() ? "List is empty" : "List is not empty") << endl;
                break;
            case 0:
                cout << "Exiting..." << endl;
                break;
            default:
                cout << "Invalid choice." << endl;
        }
    } while (choice != 0);
    
    return 0;
}""",

"Program 5: Stack using Linked List": """#include <iostream>
#include <climits> // For INT_MIN

using namespace std;

class Node {
public:
    int data;
    Node* next;
    Node(int d) {
        data = d;
        next = NULL;
    }
};

class Stack {
public:
    Node* top;
    Stack() { top = NULL; }

    bool isEmpty() {
        return top == NULL;
    }

    void push(int x) {
        Node* n = new Node(x);
        n->next = top;
        top = n;
        cout << x << " pushed to stack" << endl;
    }

    int pop() {
        if (isEmpty()) {
            cout << "Error: Stack Underflow" << endl;
            return INT_MIN;
        }
        Node* t = top;
        int x = t->data;
        top = top->next;
        delete t;
        return x;
    }

    int peek() {
        if (isEmpty()) {
            cout << "Error: Stack is Empty" << endl;
            return INT_MIN;
        }
        return top->data;
    }

    void display() {
        if (isEmpty()) {
            cout << "Stack is empty." << endl;
            return;
        }
        Node* t = top;
        cout << "Stack: TOP -> ";
        while(t) {
            cout << t->data << " -> ";
            t = t->next;
        }
        cout << "NULL" << endl;
    }
};

int main() {
    Stack s;
    int choice, val;

    do {
        cout << "\n--- Stack (Linked List) Menu ---\n";
        cout << "1. Push\n";
        cout << "2. Pop\n";
        cout << "3. Peek\n";
        cout << "4. Check if Empty\n";
        cout << "5. Display Stack\n";
        cout << "0. Exit\n";
        cout << "Enter choice: ";
        cin >> choice;

        switch (choice) {
            case 1:
                cout << "Enter value to push: ";
                cin >> val;
                s.push(val);
                break;
            case 2:
                val = s.pop();
                if (val != INT_MIN) {
                    cout << val << " popped from stack" << endl;
                }
                break;
            case 3:
                val = s.peek();
                if (val != INT_MIN) {
                    cout << "Top element is " << val << endl;
                }
                break;
            case 4:
                cout << (s.isEmpty() ? "Stack is empty" : "Stack is not empty") << endl;
                break;
            case 5:
                s.display();
                break;
            case 0:
                cout << "Exiting..." << endl;
                break;
            default:
                cout << "Invalid choice." << endl;
        }
    } while (choice != 0);

    return 0;
}""",

"Program 6: Queue using Linked List": """#include <iostream>
#include <climits> // For INT_MIN

using namespace std;

class Node {
public:
    int data;
    Node* next;
    Node(int d) {
        data = d;
        next = NULL;
    }
};

class Queue {
public:
    Node *front, *rear;
    Queue() { front = rear = NULL; }

    bool isEmpty() {
        return front == NULL;
    }

    void enqueue(int x) {
        Node* n = new Node(x);
        if (isEmpty()) {
            front = rear = n;
        } else {
            rear->next = n;
            rear = n;
        }
        cout << x << " enqueued to queue" << endl;
    }

    int dequeue() {
        if (isEmpty()) {
            cout << "Error: Queue Underflow" << endl;
            return INT_MIN;
        }
        Node* t = front;
        int x = t->data;
        front = front->next;
        if (front == NULL) {
            rear = NULL; // Queue is now empty
        }
        delete t;
        return x;
    }

    int peek() {
        if (isEmpty()) {
            cout << "Error: Queue is Empty" << endl;
            return INT_MIN;
        }
        return front->data;
    }
    
    void display() {
        if (isEmpty()) {
            cout << "Queue is empty." << endl;
            return;
        }
        Node* t = front;
        cout << "Queue: FRONT -> ";
        while(t) {
            cout << t->data << " -> ";
            t = t->next;
        }
        cout << "REAR" << endl;
    }
};

int main() {
    Queue q;
    int choice, val;

    do {
        cout << "\n--- Queue (Linked List) Menu ---\n";
        cout << "1. Enqueue\n";
        cout << "2. Dequeue\n";
        cout << "3. Peek\n";
        cout << "4. Check if Empty\n";
        cout << "5. Display Queue\n";
        cout << "0. Exit\n";
        cout << "Enter choice: ";
        cin >> choice;

        switch (choice) {
            case 1:
                cout << "Enter value to enqueue: ";
                cin >> val;
                q.enqueue(val);
                break;
            case 2:
                val = q.dequeue();
                if (val != INT_MIN) {
                    cout << val << " dequeued from queue" << endl;
                }
                break;
            case 3:
                val = q.peek();
                if (val != INT_MIN) {
                    cout << "Front element is " << val << endl;
                }
                break;
            case 4:
                cout << (q.isEmpty() ? "Queue is empty" : "Queue is not empty") << endl;
                break;
            case 5:
                q.display();
                break;
            case 0:
                cout << "Exiting..." << endl;
                break;
            default:
                cout << "Invalid choice." << endl;
        }
    } while (choice != 0);

    return 0;
}""",

"Program 7: Binary Tree Operations": """#include <iostream>
#include <algorithm> // For std::max
#include <queue> // For level-order insert

using namespace std;

class Node {
public:
    int data;
    Node *left, *right;
    Node(int d) {
        data = d;
        left = right = NULL;
    }
};

class Tree {
public:
    Node* root;
    Tree() { root = NULL; }

    // Simple level-order insert
    void insert(int x) {
        Node* n = new Node(x);
        if (root == NULL) {
            root = n;
            cout << x << " inserted as root." << endl;
            return;
        }
        queue<Node*> q;
        q.push(root);
        while (!q.empty()) {
            Node* temp = q.front();
            q.pop();
            if (temp->left == NULL) {
                temp->left = n;
                cout << x << " inserted." << endl;
                return;
            } else {
                q.push(temp->left);
            }
            if (temp->right == NULL) {
                temp->right = n;
                cout << x << " inserted." << endl;
                return;
            } else {
                q.push(temp->right);
            }
        }
    }

    void inorder(Node* r) {
        if (!r) return;
        inorder(r->left);
        cout << r->data << " ";
        inorder(r->right);
    }

    void preorder(Node* r) {
        if (!r) return;
        cout << r->data << " ";
        preorder(r->left);
        preorder(r->right);
    }

    void postorder(Node* r) {
        if (!r) return;
        postorder(r->left);
        postorder(r->right);
        cout << r->data << " ";
    }

    int countNodes(Node* r) {
        if (!r) return 0;
        return 1 + countNodes(r->left) + countNodes(r->right);
    }

    int countLeaf(Node* r) {
        if (!r) return 0;
        if (!r->left && !r->right) return 1;
        return countLeaf(r->left) + countLeaf(r->right);
    }

    int countNonLeaf(Node* r) {
        if (!r) return 0;
        if (!r->left && !r->right) return 0;
        return 1 + countNonLeaf(r->left) + countNonLeaf(r->right);
    }

    int height(Node* r) {
        if (!r) return -1; // Height of empty tree
        int l = height(r->left);
        int rt = height(r->right);
        return 1 + max(l, rt);
    }
};

int main() {
    Tree t;
    int choice, val;
    do {
        cout << "\n--- Binary Tree Menu ---\n";
        cout << "1. Insert\n";
        cout << "2. Display Inorder\n";
        cout << "3. Display Preorder\n";
        cout << "4. Display Postorder\n";
        cout << "5. Count Total Nodes\n";
        cout << "6. Count Leaf Nodes\n";
        cout << "7. Count Non-Leaf Nodes\n";
        cout << "8. Get Height\n";
        cout << "0. Exit\n";
        cout << "Enter choice: ";
        cin >> choice;

        switch (choice) {
            case 1:
                cout << "Enter value to insert: ";
                cin >> val;
                t.insert(val);
                break;
            case 2:
                cout << "Inorder: ";
                t.inorder(t.root);
                cout << endl;
                break;
            case 3:
                cout << "Preorder: ";
                t.preorder(t.root);
                cout << endl;
                break;
            case 4:
                cout << "Postorder: ";
                t.postorder(t.root);
                cout << endl;
                break;
            case 5:
                cout << "Total Nodes: " << t.countNodes(t.root) << endl;
                break;
            case 6:
                cout << "Leaf Nodes: " << t.countLeaf(t.root) << endl;
                break;
            case 7:
                cout << "Non-Leaf Nodes: " << t.countNonLeaf(t.root) << endl;
                break;
            case 8:
                cout << "Height: " << t.height(t.root) << endl;
                break;
            case 0:
                cout << "Exiting..." << endl;
                break;
            default:
                cout << "Invalid choice." << endl;
        }
    } while (choice != 0);
    return 0;
}""",

"Program 8: Binary Search Tree": """#include <iostream>
using namespace std;

class Node {
public:
    int data;
    Node *left, *right;
    Node(int d) {
        data = d;
        left = right = NULL;
    }
};

class BST {
public:
    Node* root;
    BST() { root = NULL; }

    Node* insert(Node* r, int x) {
        if (!r) {
            return new Node(x);
        }
        if (x < r->data) {
            r->left = insert(r->left, x);
        } else if (x > r->data) {
            r->right = insert(r->right, x);
        }
        // if x == r->data, do nothing (no duplicates)
        return r;
    }

    Node* minNode(Node* r) {
        if (!r) return NULL;
        while (r->left) r = r->left;
        return r;
    }

    Node* remove(Node* r, int x) {
        if (!r) return r; // Value not found

        if (x < r->data) {
            r->left = remove(r->left, x);
        } else if (x > r->data) {
            r->right = remove(r->right, x);
        } else {
            // Node with value x found
            
            // Case 1: No child
            if (!r->left && !r->right) {
                delete r;
                return NULL;
            }
            // Case 2: One child (right)
            else if (!r->left) {
                Node* t = r->right;
                delete r;
                return t;
            }
            // Case 2: One child (left)
            else if (!r->right) {
                Node* t = r->left;
                delete r;
                return t;
            }
            // Case 3: Two children
            else {
                // Find inorder successor (smallest in right subtree)
                Node* m = minNode(r->right);
                // Copy successor's data to this node
                r->data = m->data;
                // Delete the inorder successor
                r->right = remove(r->right, m->data);
            }
        }
        return r;
    }

    void inorder(Node* r) {
        if (!r) return;
        inorder(r->left);
        cout << r->data << " ";
        inorder(r->right);
    }
    
    Node* search(Node* r, int x) {
        if (!r || r->data == x) {
            return r;
        }
        if (x < r->data) {
            return search(r->left, x);
        }
        return search(r->right, x);
    }
};

int main() {
    BST b;
    int choice, val;
    
    do {
        cout << "\n--- Binary Search Tree Menu ---\n";
        cout << "1. Insert\n";
        cout << "2. Delete\n";
        cout << "3. Display Inorder\n";
        cout << "4. Search\n";
        cout << "5. Find Minimum\n";
        cout << "0. Exit\n";
        cout << "Enter choice: ";
        cin >> choice;

        switch (choice) {
            case 1:
                cout << "Enter value to insert: ";
                cin >> val;
                b.root = b.insert(b.root, val);
                cout << val << " inserted." << endl;
                break;
            case 2:
                cout << "Enter value to delete: ";
                cin >> val;
                b.root = b.remove(b.root, val);
                cout << val << " removed (if it existed)." << endl;
                break;
            case 3:
                cout << "Inorder: ";
                b.inorder(b.root);
                cout << endl;
                break;
            case 4:
                cout << "Enter value to search: ";
                cin >> val;
                if (b.search(b.root, val)) {
                    cout << val << " found in tree." << endl;
                } else {
                    cout << val << " not found in tree." << endl;
                }
                break;
            case 5: {
                Node* min = b.minNode(b.root);
                if (min) {
                    cout << "Minimum value: " << min->data << endl;
                } else {
                    cout << "Tree is empty." << endl;
                }
                break;
            }
            case 0:
                cout << "Exiting..." << endl;
                break;
            default:
                cout << "Invalid choice." << endl;
        }
    } while (choice != 0);
    
    return 0;
}""",

"Program 9: Max Heap": """#include <iostream>
#include <vector>
#include <algorithm> // for swap
#include <climits> // for INT_MIN

using namespace std;

class MaxHeap {
public:
    vector<int> arr;

    MaxHeap() {
        // We add a dummy element to make 1-based indexing easier
        arr.push_back(INT_MIN); 
    }

    int parent(int i) { return i / 2; }
    int left(int i) { return 2 * i; }
    int right(int i) { return 2 * i + 1; }
    int size() { return arr.size() - 1; }
    bool isEmpty() { return size() == 0; }

    void heapifyUp(int i) {
        while (i > 1 && arr[parent(i)] < arr[i]) {
            swap(arr[parent(i)], arr[i]);
            i = parent(i);
        }
    }

    void insert(int x) {
        arr.push_back(x);
        heapifyUp(size());
        cout << x << " inserted." << endl;
    }

    void heapifyDown(int i) {
        int l = left(i);
        int r = right(i);
        int largest = i;

        if (l <= size() && arr[l] > arr[largest]) {
            largest = l;
        }
        if (r <= size() && arr[r] > arr[largest]) {
            largest = r;
        }

        if (largest != i) {
            swap(arr[i], arr[largest]);
            heapifyDown(largest);
        }
    }

    int extractMax() {
        if (isEmpty()) {
            cout << "Error: Heap is empty" << endl;
            return INT_MIN;
        }
        int maxVal = arr[1];
        arr[1] = arr.back(); // Move last element to root
        arr.pop_back();      // Remove last element
        heapifyDown(1);      // Fix the heap
        return maxVal;
    }

    int getMax() {
        if (isEmpty()) {
            cout << "Error: Heap is empty" << endl;
            return INT_MIN;
        }
        return arr[1];
    }
    
    void display() {
        if (isEmpty()) {
            cout << "Heap is empty." << endl;
            return;
        }
        cout << "Heap: ";
        for (int i = 1; i <= size(); i++) {
            cout << arr[i] << " ";
        }
        cout << endl;
    }
};

int main() {
    MaxHeap h;
    int choice, val;

    do {
        cout << "\n--- Max Heap Menu ---\n";
        cout << "1. Insert\n";
        cout << "2. Extract Max (Delete Root)\n";
        cout << "3. Get Max\n";
        cout << "4. Display Heap\n";
        cout << "5. Check if Empty\n";
        cout << "6. Get Size\n";
        cout << "0. Exit\n";
        cout << "Enter choice: ";
        cin >> choice;

        switch (choice) {
            case 1:
                cout << "Enter value to insert: ";
                cin >> val;
                h.insert(val);
                break;
            case 2:
                val = h.extractMax();
                if (val != INT_MIN) {
                    cout << "Extracted max value: " << val << endl;
                }
                break;
            case 3:
                val = h.getMax();
                if (val != INT_MIN) {
                    cout << "Max value: " << val << endl;
                }
                break;
            case 4:
                h.display();
                break;
            case 5:
                cout << (h.isEmpty() ? "Heap is empty" : "Heap is not empty") << endl;
                break;
            case 6:
                cout << "Heap size: " << h.size() << endl;
                break;
            case 0:
                cout << "Exiting..." << endl;
                break;
            default:
                cout << "Invalid choice." << endl;
        }
    } while (choice != 0);

    return 0;
}""",

"Program 10: Sorting Algorithms": """#include <iostream>
#include <vector>
#include <algorithm> // for swap

using namespace std;

void printArray(vector<int>& a, string name) {
    cout << name << " result: ";
    for (int x : a) {
        cout << x << " ";
    }
    cout << endl;
}

// --- Selection Sort ---
void selectionSort(vector<int> a) { // Pass by value
    int n = a.size();
    for (int i = 0; i < n - 1; i++) {
        int m = i; // min index
        for (int j = i + 1; j < n; j++) {
            if (a[j] < a[m]) m = j;
        }
        swap(a[i], a[m]);
    }
    printArray(a, "Selection Sort");
}

// --- Insertion Sort ---
void insertionSort(vector<int> a) { // Pass by value
    int n = a.size();
    for (int i = 1; i < n; i++) {
        int key = a[i];
        int j = i - 1;
        while (j >= 0 && a[j] > key) {
            a[j + 1] = a[j];
            j--;
        }
        a[j + 1] = key;
    }
    printArray(a, "Insertion Sort");
}

// --- Bubble Sort ---
void bubbleSort(vector<int> a) { // Pass by value
    int n = a.size();
    for (int i = 0; i < n - 1; i++) {
        for (int j = 0; j < n - i - 1; j++) {
            if (a[j] > a[j + 1]) {
                swap(a[j], a[j + 1]);
            }
        }
    }
    printArray(a, "Bubble Sort");
}

// --- Quick Sort ---
int partition(vector<int>& a, int l, int h) {
    int p = a[h]; // pivot
    int i = l - 1;
    for (int j = l; j < h; j++) {
        if (a[j] < p) {
            i++;
            swap(a[i], a[j]);
        }
    }
    swap(a[i + 1], a[h]);
    return i + 1;
}

void quickSort(vector<int>& a, int l, int h) {
    if (l < h) {
        int p = partition(a, l, h);
        quickSort(a, l, p - 1);
        quickSort(a, p + 1, h);
    }
}

// --- Merge Sort ---
void merge(vector<int>& a, int l, int m, int h) {
    int n1 = m - l + 1;
    int n2 = h - m;
    vector<int> L(n1), R(n2);
    for (int i = 0; i < n1; i++) L[i] = a[l + i];
    for (int i = 0; i < n2; i++) R[i] = a[m + 1 + i];
    
    int i = 0, j = 0, k = l;
    while (i < n1 && j < n2) {
        if (L[i] <= R[j]) a[k++] = L[i++];
        else a[k++] = R[j++];
    }
    while (i < n1) a[k++] = L[i++];
    while (j < n2) a[k++] = R[j++];
}

void mergeSort(vector<int>& a, int l, int h) {
    if (l < h) {
        int m = l + (h - l) / 2;
        mergeSort(a, l, m);
        mergeSort(a, m + 1, h);
        merge(a, l, m, h);
    }
}

// --- Heap Sort ---
void heapify(vector<int>& a, int n, int i) {
    int largest = i, l = 2 * i + 1, r = 2 * i + 2;
    if (l < n && a[l] > a[largest]) largest = l;
    if (r < n && a[r] > a[largest]) largest = r;
    if (largest != i) {
        swap(a[i], a[largest]);
        heapify(a, n, largest);
    }
}

void heapSort(vector<int> a) { // Pass by value
    int n = a.size();
    for (int i = n / 2 - 1; i >= 0; i--) heapify(a, n, i);
    for (int i = n - 1; i >= 0; i--) {
        swap(a[0], a[i]);
        heapify(a, i, 0);
    }
    printArray(a, "Heap Sort");
}

int main() {
    cout << "Enter 5 numbers: ";
    vector<int> a(5);
    for (int i = 0; i < 5; i++) cin >> a[i];

    cout << "\nOriginal Array: ";
    for (int x : a) cout << x << " ";
    cout << "\n--- Running All Sorts ---\n";

    selectionSort(a);
    insertionSort(a);
    bubbleSort(a);
    
    vector<int> a_qs = a;
    quickSort(a_qs, 0, a.size() - 1);
    printArray(a_qs, "Quick Sort");
    
    vector<int> a_ms = a;
    mergeSort(a_ms, 0, a.size() - 1);
    printArray(a_ms, "Merge Sort");

    heapSort(a);
    
    return 0;
}"""
}

def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def create_docx(programs_list, font_name, font_size, font_color_hex):
    doc = Document()
    
    try:
        rgb = hex_to_rgb(font_color_hex)
        font_color = RGBColor(rgb[0], rgb[1], rgb[2])
    except Exception:
        font_color = RGBColor(0, 0, 0)

    for i, program in enumerate(programs_list):
        if i > 0:
            doc.add_page_break()
        
        program_number = i + 1
        full_title = f"Program {program_number}: {program['title']}"
        code = program['code']

        heading_paragraph = doc.add_heading(full_title, level=1)
        run = heading_paragraph.runs[0]
        run.font.name = 'Calibri' 
        run.font.color.rgb = RGBColor(0, 0, 0) 
        doc.add_paragraph()
        
        lines = code.split("\n")
        mid = (len(lines) + 1) // 2
        left_code = "\n".join(lines[:mid])
        right_code = "\n".join(lines[mid:])

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.columns[0].width = Pt(240)
        table.columns[1].width = Pt(240)

        cell1 = table.rows[0].cells[0]
        p1 = cell1.paragraphs[0]
        run1 = p1.add_run(left_code)
        run1.font.name = font_name
        run1.font.size = Pt(font_size)
        run1.font.color.rgb = font_color

        cell2 = table.rows[0].cells[1]
        p2 = cell2.paragraphs[0]
        run2 = p2.add_run(right_code)
        run2.font.name = font_name
        run2.font.size = Pt(font_size)
        run2.font.color.rgb = font_color
        
        doc.add_paragraph()
        doc.add_paragraph("Output:")
        doc.add_paragraph() 

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


st.set_page_config(page_title="Code to DOCX", layout="centered")
st.title("Code to DOCX Splitter ")
st.markdown("Add one or more code snippets. The app will number them and place each one in a two-column layout in a DOCX file for you to download.")

if 'programs' not in st.session_state:
    st.session_state.programs = []
if 'current_title' not in st.session_state:
    st.session_state.current_title = ""
if 'current_code' not in st.session_state:
    st.session_state.current_code = ""

def update_form_with_selection():
    selected_key = st.session_state.preloader_select
    if selected_key != "-- Select a pre-loaded program --":
        title_only = selected_key.split(":", 1)[-1].strip()
        st.session_state.current_title = title_only
        st.session_state.current_code = PRELOADED_PROGRAMS[selected_key]
    else:
        st.session_state.current_title = ""
        st.session_state.current_code = ""

def add_program_and_clear():
    if st.session_state.current_title and st.session_state.current_code:
        st.session_state.programs.append({
            "title": st.session_state.current_title,
            "code": st.session_state.current_code
        })
        st.success(f"Added 'Program {len(st.session_state.programs)}: {st.session_state.current_title}'")
        
        st.session_state.current_title = ""
        st.session_state.current_code = ""
        st.session_state.preloader_select = "-- Select a pre-loaded program --"
    
    else:
        st.error("Please provide both a title and code.")

st.header("1. Add a Program")

preloader_options = ["-- Select a pre-loaded program --"] + list(PRELOADED_PROGRAMS.keys())
st.selectbox(
    "Optional: Start with a pre-loaded program",
    options=preloader_options,
    key="preloader_select",
    on_change=update_form_with_selection
)

with st.form("add_program_form"):
    program_title = st.text_input("Program Title (e.g., 'Linear Search')", key="current_title")
    program_code = st.text_area("Paste Your Code Here", height=250, key="current_code")
    
    submitted = st.form_submit_button(
        "Add Program to List",
        on_click=add_program_and_clear 
    )

if st.session_state.programs:
    st.header("2. Set Styles & Download")

    st.subheader("Formatting Options")
    
    col1, col2 = st.columns(2)
    with col1:
        font_name = st.selectbox(
            "Select Code Font",
            ("Consolas", "Courier New", "Menlo", "Arial", "Calibri", "Times New Roman"),
            index=0
        )
        font_color = st.color_picker("Code Font Color", "#000000")
        
    with col2:
        font_size = st.number_input("Code Font Size", min_value=6, max_value=20, value=9)

    st.subheader("Document Overview")
    with st.expander("Click to see an overview of added programs"):
        if not st.session_state.programs:
            st.write("No programs added yet.")
        for i, program in enumerate(st.session_state.programs):
            st.markdown(f"**Program {i+1}: {program['title']}**")
            st.code(program['code'][:200] + "...", language=None)

    st.subheader("Download")
    try:
        docx_buffer = create_docx(st.session_state.programs, font_name, font_size, font_color)
        
        st.download_button(
            label="Download All Programs as .docx",
            data=docx_buffer,
            file_name="Split_Code_Programs.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    except Exception as e:
        st.error(f"Failed to generate DOCX file: {e}")

    if st.button("Clear All Programs", type="secondary"):
        st.session_state.programs = []
        st.session_state.current_title = ""
        st.session_state.current_code = ""
        st.session_state.preloader_select = "-- Select a pre-loaded program --"
        st.rerun()

else:
    st.info("Add a program using the form above or select a pre-loaded one to get started.")
